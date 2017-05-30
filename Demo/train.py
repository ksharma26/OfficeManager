import os;
from sklearn.feature_extraction.text import TfidfVectorizer
from pandas import DataFrame
from docx import Document

### text vectorization--go from strings to lists of numbers
vectorizer = TfidfVectorizer(sublinear_tf=True, max_df=0.5,
                             stop_words='english')

path = '.'


def read_files(path):
    for root, dirs, files in os.walk(path, topdown=False):
        for dirName in dirs:
            if (not dirName .startswith('.')):
                for filename in os.listdir(dirName):
                    if('docx' in filename):
                        document = Document(dirName + '/' + filename)
                        for para in document.paragraphs:
                            content = para.text
                    else:
                        content = open(dirName + '/' + filename, encoding="latin-1").read()
                    yield dirName, content


def build_date_frame(path):
    rows = []
    for file_name, text in read_files(path):
        rows.append({'text': text, 'class': file_name})

    data_frame = DataFrame(rows)
    return data_frame

data = build_date_frame(path)

from sklearn.feature_extraction.text import CountVectorizer
count_vectorizer = CountVectorizer()
counts = count_vectorizer.fit_transform(data['text'].values)


from sklearn.naive_bayes import MultinomialNB

classifier = MultinomialNB()
targets = data['class'].values
classifier.fit(counts, targets)

example_document = Document('C:/Users/karans/Desktop/Summers Brennan Interview.docx')
for para in example_document.paragraphs:
    example_content = para.text

examples = [example_content]
#examples = ['Social Security Number', "I'm going to attend the Linux users group tomorrow."]
example_counts = count_vectorizer.transform(examples)
predictions = classifier.predict(example_counts)
print('OfficeManager Recommendation: ',  predictions)